"""Ingestion: parse uploaded/imported content -> chunk -> embed -> immutable Asset."""

from __future__ import annotations

import hashlib
import json
import os
import uuid

from ..extensions import db
from ..models import (
    ASSET_TYPE_DOCUMENT,
    Asset,
    AssetChunk,
    AssetEmbedding,
)
from ..providers import EmbeddingProvider
from ..settings_keys import resolve_chunk_overlap, resolve_chunk_size
from ..services.settings import get_setting

# F11: per-user reindex concurrency lock (in-process). Maps user_id -> True while
# a reindex is running. Guarded by the GIL for single-process (workers=1) runtime.
_reindex_locks: dict = {}


def compute_content_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _extract_text_from_bytes(data: bytes, mime_type: str, filename: str = "") -> str:
    name = (filename or "").lower()
    if mime_type.startswith("text/") or name.endswith((".txt", ".md", ".markdown", ".csv", ".json")):
        try:
            return data.decode("utf-8", errors="replace")
        except Exception:
            return data.decode("latin-1", errors="replace")
    if name.endswith(".pdf") or mime_type == "application/pdf":
        return _extract_pdf(data)
    if name.endswith(".docx") or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(data)
    if name.endswith(".html") or mime_type in ("text/html", "application/html"):
        from bs4 import BeautifulSoup

        try:
            soup = BeautifulSoup(data.decode("utf-8", errors="replace"), "html.parser")
            for tag in soup(["script", "style"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except Exception:
            return data.decode("utf-8", errors="replace")
    # Fallback: try as text.
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        import io

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n\n".join(parts)
    except Exception:
        return ""


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except Exception:
        return ""


def chunk_text(text: str, chunk_size=None, overlap=None) -> list[str]:
    if chunk_size is None:
        chunk_size = resolve_chunk_size()
    if overlap is None:
        overlap = resolve_chunk_overlap()
    overlap = min(overlap, chunk_size - 1) if chunk_size > 1 else 0

    if not text:
        return []
    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0
    step = max(1, chunk_size - overlap)
    while start < len(words):
        slice_words = words[start : start + chunk_size]
        if not slice_words:
            break
        chunks.append(" ".join(slice_words))
        if start + chunk_size >= len(words):
            break
        start += step
    return chunks


def ingest_bytes(
    data: bytes,
    filename: str,
    mime_type: str,
    user_id: str,
    title: str | None = None,
    source_url: str | None = None,
) -> Asset:
    """Create (or return existing) immutable Asset + chunks + embeddings.

    Transactional: a failure mid-way rolls back the Asset + its chunks/embeddings.
    """
    content_hash = compute_content_hash(data)

    existing = Asset.query.filter_by(content_hash=content_hash, user_id=user_id).first()
    if existing:
        return existing

    uploads_dir = _uploads_dir()
    stored_name = f"{uuid.uuid4().hex}_{filename}"
    storage_ref = os.path.join(uploads_dir, stored_name)

    text = _extract_text_from_bytes(data, mime_type, filename)
    asset = Asset(
        content_hash=content_hash,
        content_type=ASSET_TYPE_DOCUMENT,
        storage_ref=storage_ref,
        mime_type=mime_type or "application/octet-stream",
        title=title or filename or "Untitled",
        user_id=user_id,
        status="indexing",
        metadata_json=json.dumps(
            {
                "filename": filename,
                "source_url": source_url,
                "text_length": len(text),
                "word_count": len(text.split()),
                # Binary-edit product rule (P3 #9): the ingested blob is the
                # original (immutable). Any transformation produces a NEW asset,
                # never overwrites this one.
                "is_original": True,
                "is_generated_version": False,
                "original_preserved": True,
            }
        ),
    )
    db.session.add(asset)
    db.session.flush()

    _write_file(storage_ref, data)
    _index_text(asset, text)

    asset.status = "ready"
    db.session.commit()
    return asset


def _index_text(asset: Asset, text: str):
    chunks = chunk_text(text)
    if not chunks:
        return
    # Approximate tokens (4 chars/token heuristic).
    chunk_rows = []
    for idx, chunk in enumerate(chunks):
        cr = AssetChunk(asset_id=asset.id, chunk_index=idx, text=chunk, token_count=max(1, len(chunk) // 4))
        db.session.add(cr)
        chunk_rows.append(cr)
    db.session.flush()

    model = get_setting("alfred_embedding_model", "") or _default_embedding_model()
    try:
        vectors = EmbeddingProvider.embed_batch([c.text for c in chunk_rows], model=model)
    except Exception as exc:
        # Embedding failure must not lose the Asset/chunks. Keep the Asset ready
        # (text is indexed) and mark it for later re-indexing. The file + chunks
        # remain; RAG simply won't return this asset until embeddings exist.
        asset.status = "ready"
        try:
            meta = json.loads(asset.metadata_json or "{}")
        except (ValueError, TypeError):
            meta = {}
        meta["embedding_pending"] = True
        meta["embedding_error"] = str(exc)[:300]
        asset.metadata_json = json.dumps(meta)
        db.session.flush()
        return

    for cr, vec in zip(chunk_rows, vectors):
        db.session.add(
            AssetEmbedding(
                chunk_id=cr.id,
                asset_id=asset.id,
                model=model,
                embedding=json.dumps(vec),
            )
        )
    db.session.flush()


def _default_embedding_model():
    from ..settings_keys import DEFAULT_EMBEDDING_MODEL

    return DEFAULT_EMBEDDING_MODEL


def _uploads_dir() -> str:
    from flask import current_app

    base = current_app.config.get("UPLOADS_DIR", "uploads")
    if not os.path.isabs(base):
        base = os.path.join(current_app.instance_path, base)
    uploads_dir = os.path.join(base, "assets")
    os.makedirs(uploads_dir, exist_ok=True)
    return uploads_dir


def _write_file(storage_ref: str, data: bytes):
    os.makedirs(os.path.dirname(storage_ref), exist_ok=True)
    with open(storage_ref, "wb") as fh:
        fh.write(data)


def reindex_library(user_id: str, model=None):
    """Re-embed all chunks for a user under the active embedding model.

    Atomic swap (P3 #7): new vectors are built under a temporary generation tag
    and only swapped in (delete old generation, rename temp -> active) inside a
    single transaction. Search always queries one generation, so a mid-reindex
    reader never sees a partial index.

    F11: concurrency lock — a simple in-process flag prevents two reindexes for
    the same user running at once (which would race on the temp generation and
    corrupt the index).
    """
    from .models import AssetChunk
    from sqlalchemy import text

    if _reindex_locks.get(user_id):
        raise RuntimeError("A reindex for this user is already in progress.")

    _reindex_locks[user_id] = True
    try:
        model = model or get_setting("alfred_embedding_model", "") or _default_embedding_model()
        temp_model = f"{model}#reindex"

        chunks = (
            AssetChunk.query.join(Asset)
            .filter(Asset.user_id == user_id)
            .all()
        )
        texts = [c.text for c in chunks]
        if not texts:
            return 0

        vectors = EmbeddingProvider.embed_batch(texts, model=model)

        # Build the new generation under a temp model tag (never read by search yet).
        for c, vec in zip(chunks, vectors):
            emb = AssetEmbedding(
                chunk_id=c.id, asset_id=c.asset_id, model=temp_model, embedding=json.dumps(vec)
            )
            db.session.add(emb)
        db.session.flush()

        # Atomic swap: drop the prior active generation, promote the temp generation.
        db.session.execute(
            text("DELETE FROM asset_embedding WHERE model = :model AND chunk_id IN ("
                 "SELECT id FROM asset_chunk WHERE asset_id IN "
                 "(SELECT id FROM asset WHERE user_id = :uid))"),
            {"model": model, "uid": user_id},
        )
        db.session.execute(
            text("UPDATE asset_embedding SET model = :model WHERE model = :temp"),
            {"model": model, "temp": temp_model},
        )
        db.session.commit()
        return len(chunks)
    finally:
        _reindex_locks.pop(user_id, None)
